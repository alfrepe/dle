import docx

# palabras: definicion
# def create_docx(name, content=[]):
#     doc = docx.Document()
#     for word,definition in content:
#         text = doc.add_paragraph()
#         text.add_run(word+": ").bold = True
#         text.add_run(definition)
#     doc.save(name+".docx")

#palabra
#definicion
# def create_docx2(name, content=[]):
#     doc = docx.Document()
#     for word,definition in content:
#         text = doc.add_paragraph()
#         text.add_run(word+": ").bold = True
#         text.add_run("\n")
#         text.add_run(definition)
#     doc.save(name+".docx")
#lines = [("lúbugre","Sombrío, profundamente triste"),("angosto","Estrecho o reducido")]

def write_docx(doc, word, definition):
    text = doc.add_paragraph()
    text.add_run(word+": ").bold = True
    text.add_run(definition)

def write_docx2(doc, word,definition):
    text = doc.add_paragraph()
    text.add_run(word+": ").bold = True
    text.add_run("\n")
    text.add_run(definition)

def create_docx(name, lines=[]): # lines = [(word,definition),(word,definition)...]
    doc = docx.Document()
    for line in lines:
        word, definition = line
        write_docx2(doc,word,definition)
    doc.save(name+".docx")

lines = [("lúbugre","Sombrío, profundamente triste"),("angosto","Estrecho o reducido")]
create_docx("nuevo1",lines)
