# -*- coding: utf-8 -*-
# Objetivo: leer de la base de datos aquellas palabras que tengan la columna nueva a 1, esto significará que son palabras nuevas y que aún no se ha generado un word con ellas para estudiarlas.
# No es necesario pero si recomendable que comprobemos si el archivo existe o no; porque pondremos la columna "nuevas" a null y por lo tanto no habrá palabras nuevas y no se creará el archivo docx
import sqlite3
from contextlib import closing
import os
import docx
import time

DOCX_NAME = "doc/"+time.strftime("%d-%m-%Y;%H-%M-%S")+".docx" # nomenclatura del archivo docx: día-mes-año;hora-minutos-segundos
CONSULTA_SQL = "SELECT vocabulario.palabras, vocabulario.definiciones from vocabulario where not vocabulario.nuevas is NULL"
DB_NAME = "vocabulario.db"
SET_TO_NULL_NUEVAS = "UPDATE vocabulario SET nuevas=NULL" # TODO: en lugar de poner todo el campo neuvas a null, ¿por qué no poner solo los datos que no sean null? No sé cuál será la mejor opción en términos de eficiencia.

# Formato de salida: palabra: definicion
def write_docx(doc, word, definition):
    text = doc.add_paragraph()
    text.add_run(word+": ").bold = True
    text.add_run(definition)

# palabra
# definicion
def write_docx2(doc, word, definition):
    text = doc.add_paragraph()
    text.add_run(word+": ").bold = True
    text.add_run("\n")
    text.add_run(definition)


def create_docx(name, lines=[]):  # lines = [(word,definition),(word,definition)...]
    doc = docx.Document()
    for line in lines:
        word, definition = line
        write_docx(doc, word, definition)
    doc.save(name)


def retrieve_db_data(db_name):
    if not os.path.isfile(db_name):
        raise FileNotFoundError
    data = []
    with closing(sqlite3.connect(db_name)) as conexion:
        cursor = conexion.cursor()
        cursor.execute(
            CONSULTA_SQL)
        for fila in cursor:
            data.append(fila)
    return data

# ponemos a null el campo nuevas indicando que las palabras ya han sido añadidas a un word para su posterior estudio
def set_to_null(db_name):
    with closing(sqlite3.connect(db_name)) as conexion:
        cursor = conexion.cursor()
        cursor.execute(SET_TO_NULL_NUEVAS)
        conexion.commit()


def main():
    filas = retrieve_db_data(DB_NAME)
    if filas == []:
        print("No hay palabras nuevas que estudiar, ¡lee otro libro!")
        return
    # print(len(filas))
    # for fila in filas[:2]:
    #     word, definition = fila
    #     print(word+": "+definition)
    create_docx(DOCX_NAME, filas)
    set_to_null(DB_NAME)

main()
